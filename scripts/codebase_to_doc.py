"""Codebase-to-Word document generator.

Scans a project directory recursively and produces a professional .docx
file containing every source file's contents, grouped by directory,
with a title page, table of contents, and consistent formatting.

Usage:
    python codebase_to_doc.py --path ./my_project --output docs.docx
"""

from __future__ import annotations

import argparse
import logging
import os
import sys
from collections import defaultdict
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

try:
    from tqdm import tqdm
except ImportError:
    tqdm = None

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s  %(levelname)-8s  %(message)s",
    datefmt="%H:%M:%S",
)
logger = logging.getLogger(__name__)

# ── defaults ────────────────────────────────────────────────────────────

DEFAULT_EXTENSIONS: set[str] = {
    ".py", ".js", ".jsx", ".ts", ".tsx", ".html", ".css",
    ".json", ".yaml", ".yml", ".md", ".env", ".toml",
    ".cfg", ".ini", ".sh", ".sql", ".jinja",
}

DEFAULT_SKIP_DIRS: set[str] = {
    "node_modules", "__pycache__", ".git", ".venv", "venv", "env",
    "dist", "build", ".next", ".cache", "egg-info", ".tox",
    ".mypy_cache", ".ruff_cache", ".pytest_cache", "frontend",
}

SKIP_FILES: set[str] = {
    "package-lock.json", "yarn.lock", "poetry.lock",
    ".DS_Store", "uv.lock",
}

SKIP_SUFFIXES: set[str] = {".pyc", ".pyo", ".egg"}

# ── colours / fonts ─────────────────────────────────────────────────────

CODE_BG = RGBColor(0xF5, 0xF5, 0xF5)   # light-grey code background
CODE_FONT = "Consolas"
CODE_SIZE = Pt(8)
META_COLOR = RGBColor(0x66, 0x66, 0x66)


# ── helpers ─────────────────────────────────────────────────────────────

def _human_size(nbytes: int) -> str:
    """Return a human-readable file size string."""
    for unit in ("B", "KB", "MB", "GB"):
        if nbytes < 1024:
            return f"{nbytes:.1f} {unit}"
        nbytes /= 1024
    return f"{nbytes:.1f} TB"


def _set_cell_shading(cell, hex_color: str) -> None:
    """Apply a background shading colour to a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(
        qn("w:shd"),
        {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): hex_color,
        },
    )
    shading.append(shd)


def _add_page_number(section) -> None:
    """Insert a centred page-number field into the section footer."""
    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run()
    fld_char_begin = run._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
    run._element.append(fld_char_begin)

    run2 = paragraph.add_run()
    instr = run2._element.makeelement(qn("w:instrText"), {qn("xml:space"): "preserve"})
    instr.text = " PAGE "
    run2._element.append(instr)

    run3 = paragraph.add_run()
    fld_char_end = run3._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
    run3._element.append(fld_char_end)


# ── file discovery ──────────────────────────────────────────────────────

class FileCollector:
    """Walk a directory tree and collect eligible source files."""

    def __init__(
        self,
        root: Path,
        extensions: set[str],
        skip_dirs: set[str],
    ) -> None:
        self.root = root.resolve()
        self.extensions = extensions
        self.skip_dirs = skip_dirs

    def collect(self) -> list[Path]:
        """Return a sorted list of source-file paths."""
        files: list[Path] = []
        for dirpath, dirnames, filenames in os.walk(self.root):
            # prune skipped directories in-place so os.walk doesn't descend
            dirnames[:] = [
                d for d in dirnames
                if d not in self.skip_dirs and not d.endswith(".egg-info")
            ]
            for fname in filenames:
                fp = Path(dirpath) / fname
                if fname in SKIP_FILES:
                    continue
                if fp.suffix in SKIP_SUFFIXES:
                    continue
                if fp.suffix not in self.extensions:
                    continue
                files.append(fp)

        files.sort(key=lambda p: str(p).lower())
        logger.info("Discovered %d source files under %s", len(files), self.root)
        return files


# ── document builder ────────────────────────────────────────────────────

class DocBuilder:
    """Build the Word document from collected source files."""

    def __init__(self, project_name: str, output_path: Path) -> None:
        self.project_name = project_name
        self.output_path = output_path
        self.doc = Document()
        self.total_lines = 0
        self.total_files = 0

        self._setup_styles()

    # ── styles ──────────────────────────────────────────────────────

    def _setup_styles(self) -> None:
        """Configure default styles once."""
        style = self.doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        # make sure Heading 1 / Heading 2 exist with reasonable defaults
        for level in (1, 2):
            name = f"Heading {level}"
            h = self.doc.styles[name]
            h.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # ── title page ──────────────────────────────────────────────────

    def add_title_page(self, file_count: int) -> None:
        """Add a cover page with project name, date, and file count."""
        for _ in range(6):
            self.doc.add_paragraph("")

        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(self.project_name)
        run.bold = True
        run.font.size = Pt(32)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("Codebase Documentation")
        run.font.size = Pt(18)
        run.font.color.rgb = META_COLOR

        self.doc.add_paragraph("")

        meta = self.doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = meta.add_run(
            f"Generated on {datetime.now().strftime('%B %d, %Y at %H:%M')}\n"
            f"Total files: {file_count}"
        )
        run.font.size = Pt(12)
        run.font.color.rgb = META_COLOR

        self.doc.add_page_break()

    # ── table of contents ───────────────────────────────────────────

    def add_toc(self, files: list[Path], root: Path) -> None:
        """Add a simple table-of-contents page listing every file."""
        self.doc.add_heading("Table of Contents", level=1)

        grouped: dict[str, list[Path]] = defaultdict(list)
        for fp in files:
            rel = fp.relative_to(root)
            directory = str(rel.parent) if rel.parent != Path(".") else "(root)"
            grouped[directory].append(fp)

        for directory in sorted(grouped):
            p = self.doc.add_paragraph()
            run = p.add_run(f"\n{directory}/")
            run.bold = True
            run.font.size = Pt(10)

            for fp in grouped[directory]:
                entry = self.doc.add_paragraph(style="List Bullet")
                run = entry.add_run(fp.name)
                run.font.size = Pt(9)

        self.doc.add_page_break()

    # ── file sections ───────────────────────────────────────────────

    def add_directory_heading(self, directory: str) -> None:
        """Insert a Heading 1 for a directory group."""
        self.doc.add_heading(directory, level=1)

    def add_file(self, filepath: Path, root: Path) -> bool:
        """Read a single file and append it to the document.

        Returns True on success, False if the file was skipped.
        """
        rel = filepath.relative_to(root)

        # read contents
        try:
            content = filepath.read_text(encoding="utf-8", errors="strict")
        except (UnicodeDecodeError, ValueError):
            try:
                content = filepath.read_text(encoding="latin-1")
            except Exception:
                logger.warning("Skipping binary/unreadable file: %s", rel)
                return False

        lines = content.splitlines()
        line_count = len(lines)
        file_size = filepath.stat().st_size

        # heading
        self.doc.add_heading(str(rel), level=2)

        # metadata line
        meta_p = self.doc.add_paragraph()
        run = meta_p.add_run(
            f"Size: {_human_size(file_size)}  |  Lines: {line_count}"
        )
        run.font.size = Pt(9)
        run.font.color.rgb = META_COLOR
        run.font.italic = True

        # code block inside a single-cell table with grey background
        table = self.doc.add_table(rows=1, cols=1)
        table.autofit = True
        cell = table.cell(0, 0)
        _set_cell_shading(cell, "F5F5F5")

        # clear default paragraph and write code
        cell.paragraphs[0].clear()
        code_text = content if content else "(empty file)"
        run = cell.paragraphs[0].add_run(code_text)
        run.font.name = CODE_FONT
        run.font.size = CODE_SIZE
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

        # keep paragraph spacing tight inside the cell
        fmt = cell.paragraphs[0].paragraph_format
        fmt.space_before = Pt(4)
        fmt.space_after = Pt(4)

        self.doc.add_page_break()

        self.total_lines += line_count
        self.total_files += 1
        return True

    # ── save ────────────────────────────────────────────────────────

    def save(self) -> int:
        """Write the document to disk and return its size in bytes."""
        # add page numbers to every section
        for section in self.doc.sections:
            _add_page_number(section)

        self.doc.save(str(self.output_path))
        return self.output_path.stat().st_size


# ── orchestrator ────────────────────────────────────────────────────────

def generate_document(
    root: Path,
    output: Path,
    extensions: set[str],
    skip_dirs: set[str],
) -> None:
    """Top-level function that ties collection and building together."""
    project_name = root.resolve().name.upper()

    collector = FileCollector(root, extensions, skip_dirs)
    files = collector.collect()

    if not files:
        logger.error("No matching files found under %s — nothing to do.", root)
        sys.exit(1)

    builder = DocBuilder(project_name, output)
    builder.add_title_page(len(files))
    builder.add_toc(files, root.resolve())

    # group files by directory for Heading-1 grouping
    grouped: dict[str, list[Path]] = defaultdict(list)
    for fp in files:
        rel = fp.relative_to(root.resolve())
        directory = str(rel.parent) if rel.parent != Path(".") else "(root)"
        grouped[directory].append(fp)

    iterator = sorted(grouped.items())
    if tqdm is not None:
        file_iter = tqdm(files, desc="Processing files", unit="file")
        # We still need directory headings, so we iterate grouped
        # but tick the progress bar per file.
        processed_set: set[Path] = set()

        for directory, dir_files in iterator:
            builder.add_directory_heading(directory)
            for fp in dir_files:
                builder.add_file(fp, root.resolve())
                file_iter.update(1)

        file_iter.close()
    else:
        total = len(files)
        done = 0
        for directory, dir_files in iterator:
            builder.add_directory_heading(directory)
            for fp in dir_files:
                done += 1
                print(f"  [{done}/{total}] {fp.relative_to(root.resolve())}")
                builder.add_file(fp, root.resolve())

    doc_size = builder.save()

    # summary
    print("\n" + "=" * 60)
    print("  GENERATION COMPLETE")
    print("=" * 60)
    print(f"  Files processed : {builder.total_files}")
    print(f"  Total lines     : {builder.total_lines:,}")
    print(f"  Document size   : {_human_size(doc_size)}")
    print(f"  Output          : {output.resolve()}")
    print("=" * 60)


# ── CLI ─────────────────────────────────────────────────────────────────

def parse_args(argv: list[str] | None = None) -> argparse.Namespace:
    """Parse command-line arguments."""
    parser = argparse.ArgumentParser(
        description="Generate a Word document containing your entire codebase.",
    )
    parser.add_argument(
        "--path",
        type=Path,
        default=Path("."),
        help="Root directory to scan (default: current directory)",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=Path("codebase_document.docx"),
        help="Output .docx filename (default: codebase_document.docx)",
    )
    parser.add_argument(
        "--extensions",
        type=str,
        default=None,
        help="Comma-separated extensions to include (e.g. .py,.js,.ts)",
    )
    parser.add_argument(
        "--exclude-dirs",
        type=str,
        default=None,
        help="Comma-separated additional directories to skip",
    )
    return parser.parse_args(argv)


def main(argv: list[str] | None = None) -> None:
    """Entry point."""
    args = parse_args(argv)

    root = args.path
    if not root.is_dir():
        logger.error("Path %s is not a directory.", root)
        sys.exit(1)

    extensions = DEFAULT_EXTENSIONS
    if args.extensions:
        extensions = {
            ext if ext.startswith(".") else f".{ext}"
            for ext in args.extensions.split(",")
        }

    skip_dirs = set(DEFAULT_SKIP_DIRS)
    if args.exclude_dirs:
        skip_dirs.update(args.exclude_dirs.split(","))

    generate_document(root, args.output, extensions, skip_dirs)


if __name__ == "__main__":
    main()
